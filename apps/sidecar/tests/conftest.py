"""Test fixtures for sidecar extraction tests."""

import io
from pathlib import Path

import pytest


FIXTURES_DIR = Path(__file__).parent / "fixtures"


@pytest.fixture
def sample_txt_bytes():
    return (FIXTURES_DIR / "sample.txt").read_bytes()


@pytest.fixture
def sample_txt_filename():
    return "sample.txt"


@pytest.fixture
def sample_pdf_bytes():
    return _generate_minimal_pdf("Cymek RAG Pipeline Test Document\nThis PDF contains sample text for testing extraction.")


@pytest.fixture
def sample_pdf_filename():
    return "test.pdf"


@pytest.fixture
def sample_docx_bytes():
    from docx import Document
    doc = Document()
    doc.add_paragraph("Cymek RAG Pipeline Test Document")
    doc.add_paragraph("This DOCX is used for testing extraction.")
    doc.add_paragraph("It has multiple paragraphs of content.")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@pytest.fixture
def sample_docx_filename():
    return "test.docx"


@pytest.fixture
def long_txt_bytes():
    """Generate a long text to test metadata char_count."""
    text = "Cymek platform test. " * 500
    return text.encode("utf-8")


@pytest.fixture
def long_txt_filename():
    return "long.txt"


def _generate_minimal_pdf(text: str) -> bytes:
    pages = text.split("\n")
    objects = []
    objects.append(b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj")

    num_pages = len(pages)
    page_refs = " ".join(f"{3 + i * 2} 0 R" for i in range(num_pages))
    objects.append(f"2 0 obj<</Type/Pages/Kids[{page_refs}]/Count {num_pages}>>endobj".encode())

    for i, page_text in enumerate(pages):
        content_stream = f"BT /F1 12 Tf 50 700 Td ({page_text}) Tj ET".encode()
        content_id = 3 + i * 2
        page_id = content_id + 1
        stream_len = len(content_stream)

        objects.append(
            f"{page_id} 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]"
            f"/Contents {content_id} 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj".encode()
        )
        objects.append(
            f"{content_id} 0 obj<</Length {stream_len}>>stream\n".encode()
            + content_stream
            + b"\nendstream\nendobj"
        )

    objects.append(b"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj")

    xref_offset = sum(len(o) for o in objects) + len(b"%PDF-1.4\n") + len(b"\n")

    header = b"%PDF-1.4\n"
    body = b"\n".join(objects) + b"\n"
    xref = b"xref\n0 6\n"
    offsets = [0]
    offset = len(header) + 1
    for obj in objects:
        offsets.append(offset + len(obj) + 1)
        offset += len(obj) + 1

    xref += b"0000000000 65535 f \n"
    for off in offsets[1:]:
        xref += f"{off:010d} 00000 n \n".encode()

    trailer = (
        b"trailer\n<</Size 6/Root 1 0 R>>\n"
        b"startxref\n"
        + str(len(header) + len(body) + len(xref)).encode()
        + b"\n%%EOF"
    )

    return header + body + xref + trailer
